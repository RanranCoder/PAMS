# -*- coding: utf-8 -*-
"""批量抽取党务管理系统数据库设计所需的字段结构。"""
import os, sys, io, traceback, json
from glob import glob

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

BASE = r"D:/StudyFiles/Office/党建办公室/信工党建办公室历届资料/信工党建第九届/年度部门材料汇总"

# 业务类型 -> 文件路径列表（已用 glob 确认存在）
TARGETS = {
    "扫码签到-红色讲座": [
        r"文秘部工作内容总结/“传承红色基因 赓续红色血脉”主题的讲座/“传承红色基因 赓续红色血脉”主题的讲座签到/扫码签到_信息工程学院“传承红色基因 赓续红色血脉”主题的讲座20251224165745.xlsx",
    ],
    "扫码签到-海报比赛观众": [
        r"文秘部工作内容总结/“我和我的祖国”海报设计比赛/“我和我的祖国”海报设计比赛签到表/扫码签到_信息工程学院“我和我的祖国”海报设计比赛观众签到表20251224165230.xlsx",
    ],
    "扫码签到-文明宿舍颁奖": [
        r"文秘部工作内容总结/信息工程学院文明宿舍颁奖/信息工程学院文明宿舍颁奖签到/扫码签到_信息工程学院文明宿舍颁奖签到20251224170614.xlsx",
    ],
    "签到表-换届暨动员大会": [
        r"文秘部工作内容总结/信息工程学院党建办公室第九届“换届暨动员大会”/信息工程学院党建办公室第九届“换届暨动员大会”签到表.xlsx",
    ],
    "签到表-海报比赛报名统计": [
        r"文秘部工作内容总结/“我和我的祖国”海报设计比赛/“我和我的祖国”海报设计比赛签到表/“我和我的祖国”海报比赛报名统计表.xlsx",
    ],
    "签到表-海报获奖选手": [
        r"文秘部工作内容总结/“我和我的祖国”海报设计比赛/“我和我的祖国”海报设计比赛签到表/“我和我的祖国”海报设计比赛获奖选手签到表.xlsx",
    ],
    "签到表-党员档案领取": [
        r"文秘部工作内容总结/信息工程学院党委2026届毕业生党员大会/信息工程学院党委2026届毕业生党员大会打印资料/2023级正式党员档案领取签到表.xlsx",
    ],
    "排班-控烟人员安排(第10周)": [
        r"文秘部工作内容总结/26年控烟分组/第10周禁烟劝导志愿服务活动人员安排表.xlsx",
        r"文秘部工作内容总结/26年控烟分组/单周禁烟劝导志愿服务活动人员安排表.xlsx",
    ],
    "考勤-控烟考勤docx": [
        r"文秘部工作内容总结/26年控烟分组/控烟考勤.docx",
    ],
    "排班-办公室值班表": [
        r"文秘部工作内容总结/办公室值班/26年值班表/学生组织干部办公室值班表第7周(2).xlsx",
        r"文秘部工作内容总结/办公室值班/26年值班表/二教教学楼检查安排表第3周.xlsx",
    ],
    "排班-整理党员档案安排": [
        r"文秘部工作内容总结/办公室值班/整理党员档案安排.xlsx",
        r"文秘部工作内容总结/办公室值班/毕业生成绩盖章安排.xlsx",
    ],
    "课表-干部无课表": [
        r"文秘部工作内容总结/文秘部12月26日/文秘部上学期工作内容总结/无课表/信工干部无课表/信工党建干部无课表.xlsx",
    ],
    "课表-干事无课表": [
        r"文秘部工作内容总结/文秘部12月26日/文秘部上学期工作内容总结/无课表/党建干事课表/信工党建无课表（模版）.xlsx",
    ],
    "策划书-换届暨动员大会": [
        r"组织部资料/24级F7策划书集/6月——毕业生党员大会/信息工程学院党委 2026 届毕业生党员大会策划书.docx",
    ],
    "策划书-海报设计比赛方案": [
        r"组织部资料/24级F7策划书集/10月——“我和我的祖国”海报设计比赛方案/10月——第四届“我和我的祖国”海报比赛/信息工程学院党委第四届“我和我的祖国”海报设计比赛策划书.docx",
    ],
    "策划书模版": [
        r"组织部资料/2025发的党建组织部分内事务/策划书模版.docx",
    ],
    "座位表-换届大会": [
        r"组织部资料/24级F7策划书集/6月——毕业生党员大会/2026年信息工程学院党员毕业大会座位表.xlsx",
    ],
    "座位表-海报比赛(第四届)": [
        r"组织部资料/24级F7策划书集/10月——“我和我的祖国”海报设计比赛方案/10月——第四届“我和我的祖国”海报比赛/信息工程学院党委第四届“我和我的祖国”海报设计比赛座位表.xlsx",
    ],
    "座位表-分内事务": [
        r"组织部资料/2025发的党建组织部分内事务/分内事务/座位表.xlsx",
    ],
    "座位表-红色经典诵读大赛": [
        r"组织部资料/24级F7策划书集/第十九届读书节活动之“诵读红色经典·赓续精神血脉”红色经典诵读大赛/第十九届读书节活动之“诵读红色经典·赓续精神血脉”红色经典诵读大赛座位表.xlsx",
    ],
    "议程表-毕业生党员大会": [
        r"组织部资料/24级F7策划书集/6月——毕业生党员大会/2026年毕业生党员大会活动议程表.docx",
    ],
    "议程表-海报比赛": [
        r"组织部资料/24级F7策划书集/10月——“我和我的祖国”海报设计比赛方案/10月——第四届“我和我的祖国”海报比赛/信息工程学院党委第四届“我和我的祖国”海报比赛议程表.docx",
    ],
    "议程表-分内事务": [
        r"组织部资料/2025发的党建组织部分内事务/分内事务/议程表.docx",
    ],
    "评分表-红色经典诵读决赛": [
        r"组织部资料/24级F7策划书集/第十九届读书节活动之“诵读红色经典·赓续精神血脉”红色经典诵读大赛/决赛评分表——第十九届读书节活动之“诵读红色经典·赓续精神血脉”红色经典诵读大赛活动.xlsx",
    ],
    "函调-信息收集表": [
        r"组织部资料/其他活动工作/2025.10月重点发展对象函调收集/2025下半年重点发展对象函调信息收集（信息学院）.xlsx",
    ],
    "名单-重点发展对象": [
        r"组织部资料/其他活动工作/各类名单材料——切记勿乱传/2026年上半年重点发展对象名单(信息工程学院).xlsx",
    ],
    "名单-入党积极分子推优(40期).xls": [
        r"组织部资料/其他活动工作/各类名单材料——切记勿乱传/广州城建职业学院2025年下半年“入党积极分子推优”第四十期名单.xls",
    ],
    "名单-入党积极分子通过(40期)": [
        r"组织部资料/其他活动工作/各类名单材料——切记勿乱传/40期入党积极分子通过名单.xlsx",
    ],
    "名单-入党积极分子汇总.xls": [
        r"组织部资料/其他活动工作/各类名单材料——切记勿乱传/2026年上半年入党积极分子汇总表.xls",
    ],
    "名单-学生党支部信息汇总": [
        r"组织部资料/其他活动工作/各类名单材料——切记勿乱传/党员收集/2023级学生党支部信息汇总.xlsx",
    ],
    "名单-毕业生党员组织关系转移汇总": [
        r"组织部资料/其他活动工作/各类名单材料——切记勿乱传/党员收集/广州城建职业学院毕业生党员组织关系转移汇总表（2026届）-信息工程学院.xlsx",
    ],
    "党务录入-入党申请人员信息表": [
        r"组织部资料/其他活动工作/2025下半年党务录入相关/2025入党申请人员信息表.xlsx",
    ],
    "21111工程考评表-docx": [
        r"组织部资料/其他活动工作/21111工程计划考评 收集模板/21111表格.docx",
    ],
    "21111工程-学生填写示例": [
        r"组织部资料/其他活动工作/21111工程计划考评 收集模板/学生第一党支部 李钧钦/学生第一党支部 李钧钦.docx",
    ],
    "新闻稿-换届大会": [
        r"文秘部工作内容总结/信息工程学院党建办公室第九届“换届暨动员大会”/信息工程学院党建办公室第九届“换届暨动员大会”新闻稿.docx",
    ],
    "新闻稿-红色经典诵读选拔赛": [
        r"文秘部工作内容总结/新闻稿备份/20260601学院举办红色经典诵读大赛选拔赛.docx",
    ],
    "素拓分加分表-雷锋月书本捐赠": [
        r"文秘部工作内容总结/素拓分加分表/+0.2素拓分（雷锋月书本捐赠）.xlsx",
    ],
    "素拓分加分表-禁烟知识": [
        r"文秘部工作内容总结/素拓分加分表/禁烟知识加分表.xlsx",
    ],
    "素拓分加分表-雷锋月问答": [
        r"文秘部工作内容总结/素拓分加分表/雷锋月的相关知识问答.xlsx",
    ],
    "控烟次数表": [
        r"文秘部工作内容总结/控烟/控烟次数表.xlsx",
    ],
    "摆摊-残疾人日活动": [
        r"文秘部工作内容总结/摆摊活动/“残疾人日”活动摆摊.xlsx",
        r"文秘部工作内容总结/摆摊活动/无烟活动摆摊.xlsx",
    ],
}

# 青年科技部：仅确认照片
YOUTH_PHOTO_DIR = r"青年科技部材料/罗洞志愿照"


def clean(v):
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def read_xlsx(path, max_data_rows=3):
    """返回 (sheets, report_lines)。"""
    import openpyxl
    lines = []
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    for ws in wb.worksheets:
        rows = []
        for r in ws.iter_rows():
            rows.append([clean(c.value) for c in r])
        # 去掉全空行
        rows = [r for r in rows if any(x.strip() for x in r)]
        lines.append(f"  [Sheet] {ws.title}  尺寸: {ws.max_row}行 x {ws.max_column}列 (有效非空行{len(rows)})")
        for i, r in enumerate(rows[:6]):
            lines.append(f"    R{i+1}: " + " | ".join(x if x else "·" for x in r))
        lines.append("")
    wb.close()
    return lines


def read_xls(path, max_data_rows=3):
    """老格式 .xls，用 xlrd。"""
    import xlrd
    lines = []
    wb = xlrd.open_workbook(path)
    for sheet in wb.sheets():
        lines.append(f"  [Sheet] {sheet.name}  尺寸: {sheet.nrows}行 x {sheet.ncols}列")
        for i in range(min(6, sheet.nrows)):
            vals = [clean(sheet.cell_value(i, j)) for j in range(sheet.ncols)]
            lines.append(f"    R{i+1}: " + " | ".join(x if x else "·" for x in vals))
        lines.append("")
    return lines


def docx_table_header(table, max_cell=40):
    """返回表格的所有行。"""
    out = []
    nrows = len(table.rows)
    for i, row in enumerate(table.rows):
        cells = []
        for c in row.cells:
            t = clean(c.text).replace("\n", " ")
            if len(t) > max_cell:
                t = t[:max_cell] + "…"
            cells.append(t)
        out.append(f"    TBL-R{i+1}: " + " | ".join(c if c else "·" for c in cells))
    out.append(f"    (表格共 {nrows} 行)")
    return out


def read_docx(path):
    import docx
    d = docx.Document(path)
    lines = []
    lines.append("  [段落结构] (带样式的段落)")
    for p in d.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = p.style.name if p.style else ""
        if "Heading" in style or "标题" in style:
            lines.append(f"    [{style}] {t}")
        elif t:
            lines.append(f"    {t[:120]}")
    lines.append("  [表格]")
    for tbl in d.tables:
        lines.extend(docx_table_header(tbl))
    if not d.tables:
        lines.append("    (无表格)")
    lines.append("")
    return lines


def main():
    out = []
    ok = 0
    fail = []
    for biz, rels in TARGETS.items():
        for rel in rels:
            p = os.path.join(BASE, rel)
            out.append("=" * 100)
            out.append(f"[业务类型] {biz}")
            out.append(f"[文件] {p}")
            out.append(f"[存在] {os.path.exists(p)}")
            if not os.path.exists(p):
                out.append("   !! 文件不存在，跳过")
                fail.append((p, "文件不存在"))
                out.append("")
                continue
            try:
                ext = os.path.splitext(p)[1].lower()
                if ext == ".xlsx":
                    lines = read_xlsx(p)
                elif ext == ".xls":
                    lines = read_xls(p)
                elif ext == ".docx":
                    lines = read_docx(p)
                else:
                    lines = [f"  (跳过：{ext})"]
                out.extend(lines)
                ok += 1
            except Exception as e:
                out.append(f"  !! 读取失败: {e}")
                traceback.print_exc(file=io.StringIO())
                fail.append((p, str(e)))
            out.append("")

    # 青年科技部照片确认
    photo_dir = os.path.join(BASE, YOUTH_PHOTO_DIR)
    photos = sorted(glob(os.path.join(photo_dir, "*.*")))
    out.append("=" * 100)
    out.append(f"[业务类型] 青年科技部-罗洞志愿照(照片确认)")
    out.append(f"[目录] {photo_dir}")
    out.append(f"[照片数量] {len(photos)}")
    for ph in photos[:5]:
        out.append(f"   {os.path.basename(ph)} ({os.path.getsize(ph)} bytes)")
    if len(photos) > 5:
        out.append(f"   ... 共 {len(photos)} 张，均为 jpg/png 照片")
    out.append("")

    out.append("=" * 80)
    out.append(f"成功读取: {ok}  失败: {len(fail)}")
    for p, e in fail:
        out.append(f"  失败: {p}  ->  {e}")

    result = "\n".join(out)
    # 写入结果文件
    outfile = r"D:\MyApp\PAMS\field_extract_result.txt"
    with open(outfile, "w", encoding="utf-8") as f:
        f.write(result)
    print(result)
    print("\n\n[已保存到]", outfile)


if __name__ == "__main__":
    main()
